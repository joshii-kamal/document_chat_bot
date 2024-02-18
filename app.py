import os
import openai
import json
from langchain.text_splitter import CharacterTextSplitter
from langchain.embeddings.openai import OpenAIEmbeddings
from langchain.vectorstores.faiss import FAISS
from langchain.document_loaders import TextLoader
from langchain_community.document_loaders import Docx2txtLoader
from get_config import Config
from docx import Document
from langchain_community.chat_models import ChatOpenAI
from langchain.chains import LLMChain
from langchain.prompts import ChatPromptTemplate, PromptTemplate

# from langchain.document_loaders.parsers.pdf import PDFPlumberParser

from langchain_community.document_loaders import AmazonTextractPDFLoader
import tabula



Config = Config()

os.environ["OPENAI_API_KEY"] = Config["open_ai"]["api_key"]


class Extractor:
    def __init__(self):
        self.openai_key = Config['open_ai']['api_key']
        self.openai_model = Config['open_ai']['embedding_model']
        self.openai_chat_model = Config['open_ai']['model']
        openai.api_key = self.openai_key   
        self.llm = ChatOpenAI(
            model = Config['open_ai']['model'],
            temperature = Config["open_ai"]["temperature"]
        )  


    def data_pre_processing(self,file_path):
        splited_file_path = file_path.split("/")[1].split(".")
        file_name = splited_file_path[0]
        file_ext = splited_file_path[1]

        if(file_ext == "docx"):
            doc_docx = Document(file_path)

            with open(f"{file_name}.txt", 'w') as txt_file:
                for paragraph in doc_docx.paragraphs:
                    txt_file.write(paragraph.text + '\n')

            # Extract tables
            with open(f"{file_name}.txt", 'a') as txt_file:  # Append mode to add tables after text
                for table in doc_docx.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            txt_file.write(cell.text + '\t')
                        txt_file.write('\n')
                    txt_file.write('\n')

        elif(file_ext == "pdf"):
            print("fileext::::::::::::",file_ext)
            print("file_path::::::::::::",file_path)
            tables_pdf = tabula.read_pdf(file_path, pages='all')

            print("tables_pdf::::::::::::::::::::::::",tables_pdf)
            with open(f"{file_name}pdf.txt", 'w') as txt_file:
                for table in tables_pdf:
                    txt_file.write(str(table))
                    txt_file.write('\n\n')


    def plan_embedding_with_open_ai(self, plan_file):
        try:
            
            # loader = Docx2txtLoader(plan_file)
            # documents = loader.load()

            print("plan_file:::::::::::::",plan_file)
            loader = TextLoader(plan_file)
            documents = loader.load()

            text_splitter = CharacterTextSplitter(
                separator="\n",
                chunk_size=1000,
                chunk_overlap=500
            )
            document_chunks = text_splitter.split_documents(documents)

            embeddings = OpenAIEmbeddings(openai_api_key=self.openai_key, model=self.openai_model)

            vector_db = FAISS.from_documents(document_chunks, embeddings)

            return vector_db
        except Exception as e:
            print("Error in plan_embedding function", e)

    def prompt_executor_with_open_ai(self, query, vector_db):

        try:
            context = vector_db.similarity_search(query=query, search_type="similarity")
            user_prompt_template = """

                Answer the question based only on the query:
                    context : {context}
                    query : {query}

                    Note:
                    1. You MUST extract the value and return only the value.
                    2. You MUST not give descriptive answers
                    3. You MUST only return the value or cost
                    4. If the user input asks for descriptive answer ignore note number 1,2 and 3
                    5. Make sure not to add any quotation in response returned
                    3. Simply return the output as string do not add any leading and trailing texts
            """

            prompt_template = PromptTemplate(
                template=user_prompt_template, input_variables=["context", "query"]
            )


            chain = LLMChain(llm = self.llm, prompt=prompt_template)

            input_variables = {
                "context" : context,
                "query" : query
            }

            llmRes = chain(input_variables)

            return llmRes['text']

        except Exception as e:
            print("Error in prompt_executor_with_open_ai", e)



ex = Extractor()
file_path = "./Sample_Payer_Contract.pdf"

ex.data_pre_processing(file_path)
# txt_path = "./Sample_Payer_Contract.txt"

# vector_db = ex.plan_embedding_with_open_ai(file_path)

# print("vector_db::::::::::::::::::::::::::::",vector_db)
# value = "what is the rate of MRI"
# res = ex.prompt_executor_with_open_ai(value, vector_db)
# print("res::::::::::::::::::::::::::::",res)

# with open('prompts.json', 'r') as file:
#     # Load the JSON data
#     data = json.load(file)

# output = []
# for key, value in data.items():
#     res = ex.prompt_executor_with_open_ai(value, vector_db)
#     output.append({
#         value : res
#     })
#     print("promt::::::::::::::::::::",value)

# output = json.dumps(output)
# with open("output.txt", 'a') as output_file:
#     output_file.write(output)
